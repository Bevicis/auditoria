from flask import Flask, request, jsonify
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from bs4 import BeautifulSoup
import requests
import smtplib
import os
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText

app = Flask(__name__)

RECEIVER_EMAIL = os.environ.get("RECEIVER_EMAIL", "a.utrera@bevicis.com")
SENDER_EMAIL = "a.utrera@bevicis.com"
APP_PASSWORD = os.environ.get("GMAIL_APP_PASSWORD")

def auditar_url(url):
    try:
        r = requests.get(url, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        html = r.text.lower()

        def presente(textos):
            return any(t in html for t in textos)

        resultado = []

        # Certificado HTTPS
        https_ok = url.startswith("https://")
        resultado.append((
            "Certificado HTTPS",
            "Sí" if https_ok else "No",
            "La web cuenta con certificado SSL válido." if https_ok else "No tiene HTTPS, lo cual compromete la seguridad del sitio y del usuario."
        ))

        # Aviso legal, política de privacidad y cookies
        secciones = {
            "Aviso legal": ["/aviso-legal", "aviso legal"],
            "Política de privacidad": ["/privacidad", "/politica-de-privacidad", "responsable del tratamiento", "derechos arco", "rgpd"],
            "Política de cookies": ["/cookies", "tipo de cookies", "finalidad de las cookies", "desactivar cookies"]
        }

        for nombre, palabras in secciones.items():
            cumple = presente(palabras)
            obs = f"Se ha detectado contenido o enlace relacionado con {nombre.lower()}." if cumple else f"No se ha encontrado contenido o enlace claro de {nombre.lower()}, lo cual es obligatorio."
            resultado.append((nombre, "Sí" if cumple else "No", obs))

        # Banner de cookies (gestores reconocidos)
        gestores = ["cookiebot", "consentmanager", "onetrust", "cookieyes", "iubenda", "osano"]
        banner_ok = presente(gestores)
        resultado.append((
            "Banner de cookies (gestor válido)",
            "Sí" if banner_ok else "No",
            "Se ha detectado un sistema de gestión de consentimiento reconocido." if banner_ok else "No se ha detectado Cookiebot ni un sistema similar, lo cual implica que el banner no cumple con la normativa RGPD."
        ))

        # Formularios de contacto
        forms = soup.find_all("form")
        if forms:
            checkbox_privacidad = any("checkbox" in str(f).lower() and "privacidad" in str(f).lower() for f in forms)
            marcado_por_defecto = any("checked" in str(f).lower() for f in forms if "checkbox" in str(f).lower())
            if checkbox_privacidad and not marcado_por_defecto:
                resultado.append(("Formulario con RGPD", "Sí", "Se ha detectado un checkbox de privacidad no marcado por defecto."))
            else:
                obs = "Falta checkbox de privacidad o está marcado por defecto, lo cual incumple el RGPD."
                resultado.append(("Formulario con RGPD", "No", obs))
        else:
            resultado.append(("Formulario con RGPD", "No", "No se ha detectado ningún formulario en la web."))

        # Google Analytics
        analytics = presente(["gtag", "google-analytics", "ga.js"])
        resultado.append(("Google Analytics", "Sí" if analytics else "No", "Se detectó Google Analytics." if analytics else "No se ha encontrado código de seguimiento de Google."))

        # Facebook Pixel
        pixel = presente(["fbq", "facebook.com/tr"])
        resultado.append(("Facebook Pixel", "Sí" if pixel else "No", "Se detectó el píxel de Facebook." if pixel else "No se ha detectado el píxel de Facebook."))

        return resultado
    except Exception as e:
        return [("Error al acceder a la web", "No", str(e))]

def enviar_email(asunto, cuerpo, archivos):
    msg = MIMEMultipart()
    msg["From"] = SENDER_EMAIL
    msg["To"] = RECEIVER_EMAIL
    msg["Subject"] = asunto
    msg.attach(MIMEText(cuerpo, "plain"))

    for archivo in archivos:
        with open(archivo, "rb") as f:
            part = MIMEApplication(f.read(), Name=os.path.basename(archivo))
        part["Content-Disposition"] = f'attachment; filename="{os.path.basename(archivo)}"'
        msg.attach(part)

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(SENDER_EMAIL, APP_PASSWORD)
            server.send_message(msg)
        print(f"✅ Email enviado correctamente a {RECEIVER_EMAIL}")
    except Exception as e:
        print("❌ Error al enviar el email:", str(e))

@app.route("/auditar", methods=["POST"])
def auditar():
    data = request.get_json()
    empresa = data.get("empresa", "Cliente")
    url = data.get("url", "https://example.com")
    tipo = data.get("tipo", "informativa")
    fecha = date.today().strftime("%d/%m/%Y")
    tabla = auditar_url(url)

    word_name = f"Informe_BEVI_{empresa.replace(' ', '_')}.docx"
    pdf_name = f"Informe_BEVI_{empresa.replace(' ', '_')}.pdf"

    doc = Document()
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.add_run().add_picture("logo.png", width=Inches(2.0))

    title = doc.add_paragraph("INFORME DE AUDITORÍA LEGAL WEB")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph(f"Cliente: {empresa}")
    doc.add_paragraph(f"Sitio web auditado: {url}")
    doc.add_paragraph(f"Fecha de auditoría: {fecha}")
    doc.add_paragraph("RESUMEN DEL ANÁLISIS:", style="Heading 1")

    table_doc = doc.add_table(rows=1, cols=3)
    table_doc.style = "Table Grid"
    hdr_cells = table_doc.rows[0].cells
    hdr_cells[0].text = "Aspecto Revisado"
    hdr_cells[1].text = "¿Cumple?"
    hdr_cells[2].text = "Observaciones"
    for item in tabla:
        row = table_doc.add_row().cells
        row[0].text = item[0]
        row[1].text = item[1]
        row[2].text = item[2]

    doc.add_paragraph("RECOMENDACIONES:", style="Heading 1")
    for item in tabla:
        if item[1] == "No":
            doc.add_paragraph(f"🔴 {item[2]}", style="List Bullet")

    doc.add_paragraph("Este informe ha sido elaborado por BEVICIS como parte de su servicio de auditoría legal web. Para implementar las mejoras o resolver incidencias, puede contactar con nosotros en contacto@bevicis.com o visitar www.bevicis.com.")
    doc.save(word_name)

    pdf = FPDF()
    pdf.add_page()
    pdf.image("logo.png", x=75, y=10, w=60)
    pdf.ln(30)
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "INFORME DE AUDITORÍA LEGAL WEB", ln=True, align="C")
    pdf.set_font("Arial", size=11)
    pdf.ln(10)
    pdf.multi_cell(0, 8, f"Cliente: {empresa}")
    pdf.multi_cell(0, 8, f"Sitio web auditado: {url}")
    pdf.multi_cell(0, 8, f"Fecha de auditoría: {fecha}")
    pdf.ln(5)

    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "RESUMEN DEL ANÁLISIS:", ln=True)
    pdf.set_font("Arial", size=10)
    for item in tabla:
        pdf.cell(60, 8, item[0], 1)
        pdf.cell(30, 8, item[1], 1)
        pdf.cell(95, 8, item[2][:60] + ("..." if len(item[2]) > 60 else ""), 1)
        pdf.ln()

    pdf.ln(4)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "RECOMENDACIONES:", ln=True)
    pdf.set_font("Arial", size=11)
    for item in tabla:
        if item[1] == "No":
            pdf.multi_cell(0, 8, f"- {item[2]}")

    pdf.ln(5)
    pdf.set_font("Arial", "I", 10)
    pdf.multi_cell(0, 8, "Este informe ha sido elaborado por BEVICIS como parte de su servicio de auditoría legal web. Para implementar las mejoras o resolver incidencias, puede contactar con nosotros en contacto@bevicis.com o visitar www.bevicis.com.")
    pdf.output(pdf_name)

    asunto = f"Informe de Auditoría Legal Web - {empresa}"
    cuerpo = (
        f"Hola Alejandro,\n\n"
        f"Adjunto encontrarás el informe de auditoría legal para {empresa} realizado el {fecha}.\n\n"
        f"Saludos,\nSistema automático de auditorías - BEVICIS"
    )
    enviar_email(asunto, cuerpo, [word_name, pdf_name])

    return jsonify({"status": "ok", "mensaje": "Informe generado y enviado."})

@app.route("/", methods=["GET"])
def home():
    return "API de auditorías BEVICIS operativa"

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
